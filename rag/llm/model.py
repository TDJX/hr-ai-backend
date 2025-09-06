import json
import os
from typing import Any

import pdfplumber
from langchain.schema import HumanMessage, SystemMessage
from langchain_core.embeddings import Embeddings
from langchain_core.language_models import BaseChatModel

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import docx2txt
except ImportError:
    docx2txt = None


class EmbeddingsModel:
    def __init__(self, model: Embeddings):
        self.model = model

    def get_model(self):
        return self.model


class ChatModel:
    def __init__(self, model: BaseChatModel):
        self.model = model

    def get_llm(self):
        return self.model


class ResumeParser:
    def __init__(self, chat_model: ChatModel):
        self.llm = chat_model.get_llm()
        self.resume_prompt = """
        Проанализируй текст резюме и извлеки из него структурированные данные в JSON формате.
        Верни только JSON без дополнительных комментариев.

        Формат ответа:
        {{
          "name": "Имя кандидата",
          "email": "email@example.com",
          "phone": "+7-XXX-XXX-XX-XX",
          "skills": ["навык1", "навык2", "навык3"],
          "experience": [
            {{
              "company": "Название компании",
              "position": "Должность",
              "period": "2021-2024",
              "description": "Краткое описание обязанностей"
            }}
          ],
          "total_years": 3.5,
          "education": "Образование",
          "summary": "Краткое резюме о кандидате"
        }}

        Текст резюме:
        {resume_text}
        """

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Извлекает текст из PDF файла"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                return text.strip()
        except Exception as e:
            raise Exception(f"Ошибка при чтении PDF: {str(e)}") from e

    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            print(f"[DEBUG] Extracting DOCX text from: {file_path}")

            if docx2txt:
                # Предпочитаем docx2txt для простого извлечения текста
                print("[DEBUG] Using docx2txt")
                text = docx2txt.process(file_path)
                if text:
                    print(f"[DEBUG] Extracted {len(text)} characters using docx2txt")
                    return text.strip()
                else:
                    print("[DEBUG] docx2txt returned empty text")

            if Document:
                # Используем python-docx как fallback
                print("[DEBUG] Using python-docx as fallback")
                doc = Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                print(f"[DEBUG] Extracted {len(text)} characters using python-docx")
                return text.strip()

            raise Exception(
                "Библиотеки для чтения DOCX не установлены (docx2txt или python-docx)"
            )
        except Exception as e:
            print(f"[DEBUG] DOCX extraction failed: {str(e)}")
            raise Exception(f"Ошибка при чтении DOCX: {str(e)}") from e

    def extract_text_from_doc(self, file_path: str) -> str:
        """Извлекает текст из DOC файла"""
        try:
            # Метод 1: COM автоматизация Word (самый надежный для Windows)
            import os
            if os.name == 'nt':  # Windows
                try:
                    import comtypes.client
                    print(f"[DEBUG] Trying Word COM automation for {file_path}")
                    
                    word = comtypes.client.CreateObject('Word.Application')
                    word.Visible = False
                    
                    doc = word.Documents.Open(file_path)
                    text = doc.Content.Text
                    doc.Close()
                    word.Quit()
                    
                    if text and text.strip():
                        print(f"[DEBUG] Word COM successfully extracted {len(text)} characters")
                        return text.strip()
                except Exception as e:
                    print(f"[DEBUG] Word COM failed: {e}")
            
            # Метод 2: Для .doc файлов используем python-docx
            if Document:
                try:
                    doc = Document(file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text.strip()
                except Exception:
                    # Если python-docx не может прочитать .doc, пытаемся использовать системные утилиты
                    pass

            # Попытка использовать textract (универсальная библиотека для извлечения текста)
            try:
                import textract
                print(f"[DEBUG] Using textract to process {file_path}")
                text = textract.process(file_path).decode('utf-8')
                if text and text.strip():
                    print(f"[DEBUG] textract successfully extracted {len(text)} characters")
                    return text.strip()
                else:
                    print("[DEBUG] textract returned empty text")
            except ImportError as e:
                print(f"[DEBUG] textract not available: {e}")
            except Exception as e:
                print(f"[DEBUG] textract failed: {e}")
            
            # Попытка использовать docx2txt
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                if text:
                    return text.strip()
            except ImportError:
                pass
            except Exception:
                pass
            
            # Попытка использовать oletools для старых DOC файлов
            try:
                from oletools.olevba import VBA_Parser
                from oletools import olefile
                
                if olefile.isOleFile(file_path):
                    # Это старый формат DOC, пытаемся извлечь текст
                    # Пока что возвращаем информативную ошибку
                    pass
            except ImportError:
                pass
            except Exception:
                pass

            raise Exception(
                "Не удалось извлечь текст из DOC файла ни одним из методов. "
                "Возможные причины:\n"
                "1. Microsoft Word не установлен (для COM автоматизации)\n"
                "2. Файл поврежден или не содержит текста\n"
                "3. Файл защищен паролем\n"
                "Рекомендуется конвертировать DOC в DOCX формат."
            )
        except Exception as e:
            raise Exception(f"Ошибка при чтении DOC: {str(e)}") from e

    def extract_text_from_txt(self, file_path: str) -> str:
        """Извлекает текст из TXT файла"""
        try:
            # Попробуем разные кодировки
            encodings = ["utf-8", "cp1251", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    with open(file_path, encoding=encoding) as file:
                        text = file.read()
                        return text.strip()
                except UnicodeDecodeError:
                    continue

            raise Exception("Не удалось определить кодировку текстового файла")
        except Exception as e:
            raise Exception(f"Ошибка при чтении TXT: {str(e)}") from e

    def extract_text_from_file(self, file_path: str) -> str:
        """Универсальный метод извлечения текста из файла"""
        if not os.path.exists(file_path):
            raise Exception(f"Файл не найден: {file_path}")

        # Определяем расширение файла
        _, ext = os.path.splitext(file_path.lower())

        # Добавляем отладочную информацию
        print(f"[DEBUG] Parsing file: {file_path}, detected extension: {ext}")

        if ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return self.extract_text_from_docx(file_path)
        elif ext == ".doc":
            return self.extract_text_from_doc(file_path)
        elif ext == ".txt":
            return self.extract_text_from_txt(file_path)
        else:
            raise Exception(
                f"Неподдерживаемый формат файла: {ext}. Поддерживаемые форматы: PDF, DOCX, DOC, TXT"
            )

    def parse_resume_text(self, resume_text: str) -> dict[str, Any]:
        """Парсит текст резюме через LLM"""
        try:
            messages = [
                SystemMessage(
                    content="Ты эксперт по анализу резюме. Извлекай данные точно в указанном JSON формате."
                ),
                HumanMessage(
                    content=self.resume_prompt.format(resume_text=resume_text)
                ),
            ]

            response = self.llm.invoke(messages)

            # Извлекаем JSON из ответа
            response_text = response.content.strip()

            # Пытаемся найти JSON в ответе
            if response_text.startswith("{") and response_text.endswith("}"):
                return json.loads(response_text)
            else:
                # Ищем JSON внутри текста
                start = response_text.find("{")
                end = response_text.rfind("}") + 1
                if start != -1 and end > start:
                    json_str = response_text[start:end]
                    return json.loads(json_str)
                else:
                    raise ValueError("JSON не найден в ответе LLM")

        except json.JSONDecodeError as e:
            raise Exception(f"Ошибка парсинга JSON из ответа LLM: {str(e)}") from e
        except Exception as e:
            raise Exception(f"Ошибка при обращении к LLM: {str(e)}") from e

    def parse_resume_from_file(self, file_path: str) -> dict[str, Any]:
        """Полный цикл парсинга резюме из файла"""
        # Шаг 1: Извлекаем текст из файла (поддерживаем PDF, DOCX, DOC, TXT)
        resume_text = self.extract_text_from_file(file_path)

        if not resume_text:
            raise Exception("Не удалось извлечь текст из файла")

        # Шаг 2: Парсим через LLM
        return self.parse_resume_text(resume_text)
