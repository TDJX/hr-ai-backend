import io
import json
import logging
from pathlib import Path
from typing import Any, Dict

logger = logging.getLogger(__name__)


class VacancyParserService:
    """Сервис для парсинга вакансий из файлов различных форматов"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.rtf', '.txt']

    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """
        Извлекает текст из файла в зависимости от его формата
        
        Args:
            file_content: Содержимое файла в байтах
            filename: Имя файла для определения формата
            
        Returns:
            str: Извлеченный текст
        """
        file_extension = Path(filename).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_content)
            elif file_extension == '.rtf':
                return self._extract_from_rtf(file_content)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_content)
            else:
                raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")
                
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из файла {filename}: {str(e)}")
            raise

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Извлекает текст из PDF файла"""
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
            return text.strip()
            
        except ImportError:
            # Fallback to pdfplumber if PyPDF2 doesn't work well
            try:
                import pdfplumber
                
                pdf_file = io.BytesIO(file_content)
                text = ""
                
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            
                return text.strip()
                
            except ImportError:
                raise ImportError("Требуется установить PyPDF2 или pdfplumber: pip install PyPDF2 pdfplumber")

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            import docx
            
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
                    
            return text.strip()
            
        except ImportError:
            raise ImportError("Требуется установить python-docx: pip install python-docx")

    def _extract_from_rtf(self, file_content: bytes) -> str:
        """Извлекает текст из RTF файла"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            rtf_content = file_content.decode('utf-8', errors='ignore')
            text = rtf_to_text(rtf_content)
            
            return text.strip()
            
        except ImportError:
            raise ImportError("Требуется установить striprtf: pip install striprtf")
        except Exception as e:
            # Альтернативный метод через pyth
            try:
                from pyth.plugins.rtf15.reader import Rtf15Reader
                from pyth.plugins.plaintext.writer import PlaintextWriter
                
                doc = Rtf15Reader.read(io.BytesIO(file_content))
                text = PlaintextWriter.write(doc).getvalue()
                
                return text.strip()
                
            except ImportError:
                raise ImportError("Требуется установить striprtf или pyth: pip install striprtf pyth")

    def _extract_from_txt(self, file_content: bytes) -> str:
        """Извлекает текст из TXT файла"""
        try:
            # Пробуем различные кодировки
            encodings = ['utf-8', 'windows-1251', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
                    
            # Если все кодировки не подошли, используем errors='ignore'
            text = file_content.decode('utf-8', errors='ignore')
            return text.strip()
            
        except Exception as e:
            logger.error(f"Ошибка при чтении txt файла: {str(e)}")
            raise

    async def parse_vacancy_with_ai(self, raw_text: str) -> Dict[str, Any]:
        """
        Парсит текст вакансии с помощью AI для извлечения структурированной информации
        
        Args:
            raw_text: Сырой текст вакансии
            
        Returns:
            Dict с полями для модели Vacancy
        """
        from rag.settings import settings
        
        if not settings.openai_api_key:
            raise ValueError("OpenAI API ключ не настроен")
            
        try:
            import openai
            
            openai.api_key = settings.openai_api_key
            
            parsing_prompt = f"""
Проанализируй текст вакансии и извлеки из него структурированную информацию.

ТЕКСТ ВАКАНСИИ:
{raw_text}

ЗАДАЧА:
Извлеки следующие поля для вакансии:

1. title - название позиции (строка)
2. description - описание вакансии (полное описание обязанностей, требований)
3. key_skills - ключевые навыки через запятую (строка)
4. employment_type - тип занятости: "full", "part", "project", "volunteer", "probation"
5. experience - опыт работы: "noExperience", "between1And3", "between3And6", "moreThan6"
6. schedule - график работы: "fullDay", "shift", "flexible", "remote", "flyInFlyOut"
7. salary_from - зарплата от (число или null)
8. salary_to - зарплата до (число или null)  
9. salary_currency - валюта (строка, по умолчанию "RUR")
10. company_name - название компании (строка)
11. company_description - описание компании (строка или null)
12. area_name - город/регион (строка)
13. address - адрес (строка или null)
14. professional_roles - профессиональные роли (строка или null)
15. contacts_name - контактное лицо (строка или null)
16. contacts_email - email для связи (строка или null)
17. contacts_phone - телефон для связи (строка или null)

ПРАВИЛА:
- Если информация не найдена, ставь null для необязательных полей
- Для обязательных полей используй разумные значения по умолчанию
- Зарплату указывай в рублях, конвертируй если нужно
- Опыт определяй по годам: 0-1 = noExperience, 1-3 = between1And3, 3-6 = between3And6, 6+ = moreThan6
- График работы определяй по описанию: офис = fullDay, удаленка = remote, гибкий = flexible

ФОРМАТИРОВАНИЕ ТЕКСТА:
- Если в тексте есть списки (обязанности, требования, навыки), форматируй их с переносами строк
- Используй символ \n для переноса строки между пунктами списка
- Пример: "Обязанности:\nВедение переговоров\nПодготовка документов\nОбучение персонала"
- Для ключевых навыков разделяй запятыми, но если их много - используй переносы строк
- В описании компании тоже используй переносы для лучшей читаемости

ОТВЕТЬ СТРОГО В JSON ФОРМАТЕ с указанными полями:
"""

            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": parsing_prompt}],
                response_format={"type": "json_object"},
            )

            parsed_data = json.loads(response.choices[0].message.content)
            
            # Валидируем и обрабатываем данные
            return self._validate_parsed_data(parsed_data)
            
        except Exception as e:
            logger.error(f"Ошибка при парсинге вакансии через AI: {str(e)}")
            raise

    def _validate_parsed_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Валидирует и очищает спарсенные данные"""
        from app.models.vacancy import EmploymentType, Experience, Schedule
        
        # Обязательные поля с дефолтными значениями
        validated_data = {
            'title': data.get('title', 'Название не указано'),
            'description': data.get('description', 'Описание не указано'),
            'key_skills': data.get('key_skills'),
            'employment_type': self._validate_enum(
                data.get('employment_type'), 
                EmploymentType, 
                EmploymentType.FULL_TIME
            ),
            'experience': self._validate_enum(
                data.get('experience'), 
                Experience, 
                Experience.BETWEEN_1_AND_3
            ),
            'schedule': self._validate_enum(
                data.get('schedule'), 
                Schedule, 
                Schedule.FULL_DAY
            ),
            'company_name': data.get('company_name'),
            'area_name': data.get('area_name'),
        }
        
        # Необязательные поля
        optional_fields = [
            'salary_from', 'salary_to', 'salary_currency', 'company_description',
            'address', 'professional_roles', 'contacts_name', 'contacts_email', 'contacts_phone'
        ]
        
        for field in optional_fields:
            value = data.get(field)
            if value and value != "null":
                validated_data[field] = value
        
        # Специальная обработка зарплаты
        if data.get('salary_from'):
            try:
                validated_data['salary_from'] = int(data['salary_from'])
            except (ValueError, TypeError):
                pass
                
        if data.get('salary_to'):
            try:
                validated_data['salary_to'] = int(data['salary_to'])
            except (ValueError, TypeError):
                pass
        
        # Валюта по умолчанию
        validated_data['salary_currency'] = data.get('salary_currency', 'RUR')
        
        return validated_data

    def _validate_enum(self, value: str, enum_class, default_value):
        """Валидирует значение enum"""
        if not value:
            return default_value
            
        # Проверяем, есть ли такое значение в enum
        try:
            return enum_class(value)
        except ValueError:
            logger.warning(f"Неизвестное значение {value} для {enum_class.__name__}, используем {default_value}")
            return default_value


# Экземпляр сервиса
vacancy_parser_service = VacancyParserService()